from qdrant_client import QdrantClient
from sentence_transformers import SentenceTransformer
from docx import Document
import together  # ✅ Use Together.ai instead of openai
import os

# Qdrant config
QDRANT_API_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJhY2Nlc3MiOiJtIn0.JZHaC8ey7NaWyKJmWDkOdyvIhobFMhsNhQaQzU1MH7U"
QDRANT_HOST = "https://b08b781e-60b0-4f1e-a385-06a4e7c894bd.eu-west-1-0.aws.cloud.qdrant.io"
COLLECTION_NAME = "rfp_docs"

# Together.ai config
together.api_key = "845270fb939240555c4c4b93737e04a4c14e05ee90b636cbafb35fbf03e950fb"  # 🔑 Replace with your Together.ai key

client = QdrantClient(url=QDRANT_HOST, api_key=QDRANT_API_KEY)
embedder = SentenceTransformer("all-MiniLM-L6-v2")

def query_qdrant(question, top_k=5):
    question_vec = embedder.encode(question).tolist()
    hits = client.search(
        collection_name=COLLECTION_NAME,
        query_vector=question_vec,
        limit=top_k
    )
    return [hit.payload["text"] for hit in hits]

def generate_answer(question):
    context_chunks = query_qdrant(question)
    context = "\n".join(context_chunks)

    prompt = f"""
Answer the following question based on the context below.

Context:
{context}

Question:
{question}

Answer:
"""

    response = together.Complete.create(
        prompt=prompt,
        model="mistralai/Mixtral-8x7B-Instruct-v0.1",
        max_tokens=500,
        temperature=0.7,
    )

    return response['choices'][0]['text'].strip()


def write_to_docx(doc, question, answer):
    doc.add_heading("Question", level=2)
    doc.add_paragraph(question)
    doc.add_heading("Answer", level=2)
    doc.add_paragraph(answer)


    

def run_qa_pipeline(question_file="data/questions.txt", output_path="output.docx"):
    doc = Document()
    doc.add_heading("Q&A Output", level=1)

    with open(question_file, "r", encoding="utf-8") as f:
        questions = [line.strip() for line in f if line.strip()]
    
    for question in questions:
        answer = generate_answer(question)
        write_to_docx(doc, question, answer)

    doc.save(output_path)
    print(f"[Docx] All answers saved to '{output_path}'")




__all__ = ["run_qa_pipeline"]